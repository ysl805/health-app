import os
import json
import base64
import httpx
from dotenv import load_dotenv

# Load .env from project root
load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env'))
from datetime import datetime
from typing import List, Optional
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from sqlalchemy import func, distinct
from sqlalchemy.orm import Session
from pydantic import BaseModel

from database import engine, get_db, Base, migrate_db
from models import User, KnowledgeBase, KnowledgeDocument, KnowledgeBaseBinding, Consultation, TongueImage, LocalCase
from auth import (
    verify_password, get_password_hash, create_access_token,
    get_current_user, require_role, can_manage, SUBORDINATE_ROLES, ROLE_HIERARCHY
)
from ai_service import call_agnes_ai, analyze_tongue_image, stream_agnes_ai

# Create tables
Base.metadata.create_all(bind=engine)
migrate_db()

class UTF8JSONResponse(JSONResponse):
    media_type = "application/json; charset=utf-8"

app = FastAPI(title="个体化精准养生APP", version="1.0.0", default_response_class=UTF8JSONResponse)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== Frontend Static Files ====================
from fastapi.staticfiles import StaticFiles
import pathlib

FRONTEND_DIR = pathlib.Path(__file__).parent.parent / "frontend"
if FRONTEND_DIR.exists():
    # Mount frontend at root (must be last, after all API routes)
    pass  # Will mount after all routes are defined


UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ==================== Pydantic Schemas ====================

class LoginRequest(BaseModel):
    username: str
    password: str

class UserCreate(BaseModel):
    username: str
    password: str
    role: str
    real_name: str = ""
    province: str = ""
    region: str = ""
    valid_from: Optional[str] = None  # ISO datetime string
    valid_until: Optional[str] = None
    knowledge_base_ids: Optional[List[int]] = None  # 绑定知识库ID列表

class UserUpdate(BaseModel):
    real_name: Optional[str] = None
    province: Optional[str] = None
    region: Optional[str] = None
    is_active: Optional[bool] = None
    password: Optional[str] = None
    valid_from: Optional[str] = None
    valid_until: Optional[str] = None

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

class KnowledgeBaseCreate(BaseModel):
    name: str
    description: str = ""
    prompt_config: str = ""
    negative_prompt: str = ""  # 禁止回复内容

class ConsultationCreate(BaseModel):
    question: str
    knowledge_base_id: Optional[int] = None
    tongue_image_base64: Optional[str] = None

class KnowledgeBaseBindRequest(BaseModel):
    knowledge_base_id: int
    user_ids: List[int]


# ==================== Init Super Admin ====================

def init_super_admin():
    migrate_db()
    from database import SessionLocal
    db = SessionLocal()
    try:
        admin = db.query(User).filter(User.username == "superadmin").first()
        if not admin:
            admin = User(
                username="superadmin",
                hashed_password=get_password_hash("admin123456"),
                role="super_admin",
                real_name="超级管理员",
                is_active=True,
            )
            db.add(admin)
            db.commit()
            print("[OK] Super admin created: superadmin / admin123456")
        else:
            print("[INFO] Super admin already exists")
    finally:
        db.close()


# ==================== Auth APIs ====================

@app.post("/api/auth/login")
async def login(request: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == request.username).first()
    if not user or not verify_password(request.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="用户名或密码错误")
    if not user.is_active:
        raise HTTPException(status_code=403, detail="账号已被禁用")
    # Check account validity period
    now = datetime.utcnow()
    if user.valid_from and now < user.valid_from:
        raise HTTPException(status_code=403, detail="账号尚未生效，有效期开始时间为 {}".format(user.valid_from.strftime("%Y-%m-%d %H:%M")))
    if user.valid_until and now > user.valid_until:
        # Auto-disable expired accounts
        user.is_active = False
        db.commit()
        raise HTTPException(status_code=403, detail="账号已过有效期")

    token = create_access_token(data={"sub": user.username, "role": user.role})
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "username": user.username,
            "role": user.role,
            "real_name": user.real_name,
            "province": user.province,
            "region": user.region,
        },
    }


@app.get("/api/auth/me")
async def get_me(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Auto-check expiry on every request
    now = datetime.utcnow()
    if current_user.valid_until and now > current_user.valid_until and current_user.is_active:
        current_user.is_active = False
        db.commit()
    return {
        "id": current_user.id,
        "username": current_user.username,
        "role": current_user.role,
        "real_name": current_user.real_name,
        "province": current_user.province,
        "region": current_user.region,
        "valid_from": current_user.valid_from.isoformat() + "+08:00" if current_user.valid_from else None,
        "valid_until": current_user.valid_until.isoformat() + "+08:00" if current_user.valid_until else None,
    }

@app.post("/api/auth/change-password")
async def change_password(
    req: ChangePasswordRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not verify_password(req.old_password, current_user.hashed_password):
        raise HTTPException(status_code=400, detail="原密码错误")
    if len(req.new_password) < 6:
        raise HTTPException(status_code=400, detail="新密码长度不能少于6位")
    current_user.hashed_password = get_password_hash(req.new_password)
    db.commit()
    return {"message": "密码修改成功"}


# ==================== User Management APIs ====================

@app.post("/api/users")
async def create_user(
    user_data: UserCreate,
    current_user: User = Depends(require_role("super_admin", "admin", "province_manager", "region_manager", "business_manager")),
    db: Session = Depends(get_db),
):
    # Check if current user can create this role
    if not can_manage(current_user.role, user_data.role):
        raise HTTPException(status_code=403, detail=f"您无权创建 {user_data.role} 角色的用户")

    # Check username uniqueness
    if db.query(User).filter(User.username == user_data.username).first():
        raise HTTPException(status_code=400, detail="用户名已存在")

    new_user = User(
        username=user_data.username,
        hashed_password=get_password_hash(user_data.password),
        role=user_data.role,
        real_name=user_data.real_name,
        province=user_data.province,
        region=user_data.region,
        valid_from=datetime.fromisoformat(user_data.valid_from) if user_data.valid_from else None,
        valid_until=datetime.fromisoformat(user_data.valid_until) if user_data.valid_until else None,
        created_by=current_user.id,
        is_active=True,
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)

    # 绑定知识库
    if user_data.knowledge_base_ids:
        for kb_id in user_data.knowledge_base_ids:
            binding = KnowledgeBaseBinding(user_id=new_user.id, knowledge_base_id=kb_id)
            db.add(binding)
        db.commit()

    return {"id": new_user.id, "username": new_user.username, "role": new_user.role}


@app.get("/api/users")
async def list_users(
    role: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(User)

    # Non-super_admin can only see subordinate users
    if current_user.role != "super_admin":
        allowed_roles = SUBORDINATE_ROLES.get(current_user.role, [])
        query = query.filter(User.role.in_(allowed_roles))
        # Further filter: only users created by this user's chain
        if current_user.role in ("province_manager", "region_manager", "business_manager"):
            query = query.filter(
                (User.province == current_user.province) if current_user.province else (User.id == -1)
            )

    if role:
        query = query.filter(User.role == role)

    users = query.all()
    # Get knowledge base bindings for all users
    user_ids = [u.id for u in users]
    bindings = db.query(KnowledgeBaseBinding).filter(KnowledgeBaseBinding.user_id.in_(user_ids)).all()
    user_kb_map = {}
    for b in bindings:
        if b.user_id not in user_kb_map:
            user_kb_map[b.user_id] = []
        user_kb_map[b.user_id].append(b.knowledge_base_id)
    return [{
        "id": u.id, "username": u.username, "role": u.role,
        "real_name": u.real_name, "province": u.province,
        "region": u.region, "is_active": u.is_active,
        "valid_from": u.valid_from.isoformat() + "+08:00" if u.valid_from else None,
        "valid_until": u.valid_until.isoformat() + "+08:00" if u.valid_until else None,
        "created_at": u.created_at.isoformat() + "+08:00" if u.created_at else None,
        "knowledge_base_ids": user_kb_map.get(u.id, []),
    } for u in users]


@app.put("/api/users/{user_id}")
async def update_user(
    user_id: int,
    user_data: UserUpdate,
    current_user: User = Depends(require_role("super_admin", "admin", "province_manager", "region_manager", "business_manager")),
    db: Session = Depends(get_db),
):
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(status_code=404, detail="用户不存在")
    if not can_manage(current_user.role, target.role):
        raise HTTPException(status_code=403, detail="权限不足")

    if user_data.real_name is not None:
        target.real_name = user_data.real_name
    if user_data.province is not None:
        target.province = user_data.province
    if user_data.region is not None:
        target.region = user_data.region
    if user_data.is_active is not None:
        target.is_active = user_data.is_active
    if user_data.password:
        target.hashed_password = get_password_hash(user_data.password)
    if user_data.valid_from is not None:
        target.valid_from = datetime.fromisoformat(user_data.valid_from) if user_data.valid_from else None
    if user_data.valid_until is not None:
        target.valid_until = datetime.fromisoformat(user_data.valid_until) if user_data.valid_until else None

    db.commit()
    return {"message": "更新成功"}


@app.delete("/api/users/{user_id}")
async def delete_user(
    user_id: int,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(status_code=404, detail="用户不存在")
    target.is_active = False
    db.commit()
    return {"message": "用户已禁用"}


# ==================== Knowledge Base APIs ====================

@app.post("/api/knowledge-bases")
async def create_knowledge_base(
    kb_data: KnowledgeBaseCreate,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    kb = KnowledgeBase(
        name=kb_data.name,
        description=kb_data.description,
        prompt_config=kb_data.prompt_config or "",
        negative_prompt=kb_data.negative_prompt or "",
        created_by=current_user.id,
    )
    db.add(kb)
    db.commit()
    db.refresh(kb)
    return {"id": kb.id, "name": kb.name}


@app.get("/api/knowledge-bases")
async def list_knowledge_bases(
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    kbs = db.query(KnowledgeBase).all()
    result = []
    for kb in kbs:
        doc_count = db.query(KnowledgeDocument).filter(KnowledgeDocument.knowledge_base_id == kb.id).count()
        binding_count = db.query(KnowledgeBaseBinding).filter(KnowledgeBaseBinding.knowledge_base_id == kb.id).count()
        result.append({
            "id": kb.id, "name": kb.name, "description": kb.description,
            "prompt_config": kb.prompt_config or "",
            "negative_prompt": kb.negative_prompt or "",
            "doc_count": doc_count, "binding_count": binding_count,
            "created_at": kb.created_at.isoformat() + "+08:00" if kb.created_at else None,
        })
    return result


@app.put("/api/knowledge-bases/{kb_id}")
async def update_knowledge_base(
    kb_id: int,
    kb_data: KnowledgeBaseCreate,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")
    kb.name = kb_data.name
    kb.description = kb_data.description
    kb.prompt_config = kb_data.prompt_config or ""
    kb.negative_prompt = kb_data.negative_prompt or ""
    db.commit()
    return {"message": "更新成功"}


@app.delete("/api/knowledge-bases/{kb_id}")
async def delete_knowledge_base(
    kb_id: int,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")
    # Delete related documents and bindings
    db.query(KnowledgeDocument).filter(KnowledgeDocument.knowledge_base_id == kb_id).delete()
    db.query(KnowledgeBaseBinding).filter(KnowledgeBaseBinding.knowledge_base_id == kb_id).delete()
    db.delete(kb)
    db.commit()
    return {"message": "知识库已删除"}


@app.post("/api/knowledge-bases/{kb_id}/documents")
async def upload_document(
    kb_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    # Save file
    file_content = await file.read()
    file_path = os.path.join(UPLOAD_DIR, f"kb_{kb_id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{file.filename}")
    with open(file_path, "wb") as f:
        f.write(file_content)

    # Parse content based on file type
    parsed_content = ""
    ext = os.path.splitext(file.filename)[1].lower()

    if ext in (".txt", ".md", ".csv"):
        try:
            parsed_content = file_content.decode("utf-8")
        except:
            parsed_content = file_content.decode("gbk", errors="replace")
    elif ext == ".json":
        try:
            data = json.loads(file_content.decode("utf-8"))
            parsed_content = json.dumps(data, ensure_ascii=False, indent=2)
        except:
            parsed_content = str(file_content)
    elif ext == ".pdf":
        try:
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            pages = []
            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"--- 第{i+1}页 ---\n{text.strip()}")
            parsed_content = "\n\n".join(pages) if pages else ""
        except Exception as e:
            parsed_content = f"[PDF解析失败: {str(e)}]"
    elif ext in (".docx", ".doc"):
        try:
            import docx as docx_lib
            import io
            doc = docx_lib.Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip("| "):
                        tables_text.append(row_text)
            all_text = paragraphs
            if tables_text:
                all_text.append("\n--- 表格内容 ---")
                all_text.extend(tables_text)
            parsed_content = "\n".join(all_text) if all_text else ""
        except Exception as e:
            parsed_content = f"[Word解析失败: {str(e)}]"
    elif ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"):
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(file_content))
            # Store image info as parsed content (actual image analysis done via AI)
            info = f"[图片文件] 文件名: {file.filename}, 尺寸: {img.size[0]}x{img.size[1]}, 格式: {img.format or ext}"
            # Save base64 for AI analysis
            import base64
            buf = io.BytesIO()
            img.save(buf, format=img.format or "PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            parsed_content = info + "\n[base64]" + b64
        except Exception as e:
            parsed_content = f"[图片解析失败: {str(e)}]"
    else:
        parsed_content = f"[不支持的格式: {ext}]"

    doc = KnowledgeDocument(
        knowledge_base_id=kb_id,
        filename=file.filename,
        file_path=file_path,
        file_type=os.path.splitext(file.filename)[1],
        parsed_content=parsed_content,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {"id": doc.id, "filename": doc.filename, "parsed": bool(parsed_content)}


@app.get("/api/knowledge-bases/{kb_id}/documents")
async def list_documents(
    kb_id: int,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    docs = db.query(KnowledgeDocument).filter(KnowledgeDocument.knowledge_base_id == kb_id).all()
    return [{
        "id": d.id, "filename": d.filename, "file_type": d.file_type,
        "parsed": bool(d.parsed_content),
        "parsed_content": d.parsed_content or "",
        "content_length": len(d.parsed_content) if d.parsed_content else 0,
        "created_at": d.created_at.isoformat() + "+08:00" if d.created_at else None,
    } for d in docs]


@app.delete("/api/knowledge-bases/{kb_id}/documents/{doc_id}")
async def delete_document(
    kb_id: int,
    doc_id: int,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.id == doc_id,
        KnowledgeDocument.knowledge_base_id == kb_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    db.delete(doc)
    db.commit()
    return {"message": "文档已删除"}


# ==================== Knowledge Base Binding APIs ====================

@app.post("/api/knowledge-bases/bind")
async def bind_knowledge_base(
    bind_data: KnowledgeBaseBindRequest,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == bind_data.knowledge_base_id).first()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    # Remove existing bindings for this KB
    db.query(KnowledgeBaseBinding).filter(
        KnowledgeBaseBinding.knowledge_base_id == bind_data.knowledge_base_id
    ).delete()

    for user_id in bind_data.user_ids:
        user = db.query(User).filter(User.id == user_id).first()
        if user:
            binding = KnowledgeBaseBinding(
                knowledge_base_id=bind_data.knowledge_base_id,
                user_id=user_id,
            )
            db.add(binding)

    db.commit()
    return {"message": "知识库绑定成功"}


@app.get("/api/knowledge-bases/{kb_id}/bindings")
async def get_bindings(
    kb_id: int,
    current_user: User = Depends(require_role("super_admin")),
    db: Session = Depends(get_db),
):
    bindings = db.query(KnowledgeBaseBinding).filter(KnowledgeBaseBinding.knowledge_base_id == kb_id).all()
    return [{
        "id": b.id, "user_id": b.user_id,
        "username": b.user.username if b.user else "",
        "role": b.user.role if b.user else "",
    } for b in bindings]


# ==================== User's accessible knowledge bases ====================

@app.get("/api/my/knowledge-bases")
async def my_knowledge_bases(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # Super admin sees all knowledge bases
    if current_user.role == "super_admin":
        kbs = db.query(KnowledgeBase).all()
        return [{"id": kb.id, "name": kb.name, "description": kb.description} for kb in kbs]
    # Other roles: only bound knowledge bases
    bindings = db.query(KnowledgeBaseBinding).filter(KnowledgeBaseBinding.user_id == current_user.id).all()
    result = []
    for b in bindings:
        kb = b.knowledge_base
        result.append({
            "id": kb.id, "name": kb.name, "description": kb.description,
        })
    return result


# ==================== AI Consultation APIs ====================


# ==================== Tongue Image Analysis ====================

async def analyze_tongue_image(base64_image: str, knowledge_content: list = None) -> dict:
    """
    调用 AI API 分析舌象图片
    返回: {"tongue_analysis": "分析结果"}
    """
    try:
        # 构建系统提示词
        system_prompt = """你是专业的中医舌象辨析专家。请仔细分析用户上传的舌象图片，从以下几个方面进行辨析：

1. **舌质（舌色）**：淡白、淡红、红、绛红、青紫等
2. **舌苔**：颜色（白、黄、灰、黑）、厚薄（薄、厚、少苔、无苔）、润燥（润、燥、滑、糙）
3. **舌形**：老嫩、胖瘦、齿痕、裂纹、芒刺
4. **舌态**：痿软、强硬、歪斜、颤动、吐弄
5. **舌下络脉**：颜色、形态、曲张情况

请按照以下格式输出：
◆ 舌象辨析结果 ◆
【舌质】...(颜色、光泽)
【舌苔】...(颜色、厚薄、润燥)
【舌形】...(形态、特殊标记)
【舌态】...(动态特征)
【舌下络脉】...(若可见)
【综合判断】...(中医病机分析)

注意：
- 描述要客观、专业
- 避免主观臆断
- 如果图片不清晰，请说明
- 结合中医理论进行分析"""

        # 如果有知识库内容，追加到提示词
        if knowledge_content:
            system_prompt += "\n\n【参考资料】\n" + "\n".join(knowledge_content)

        # 构建用户消息（包含图片）
        if "base64," in base64_image:
            image_url = base64_image
        else:
            image_url = f"data:image/jpeg;base64,{base64_image}"

        messages = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": image_url}},
                    {"type": "text", "text": "请分析这个舌象图片，给出专业的中医舌象辨析结果。"}
                ]
            }
        ]

        # 调用 Agnes AI API
        AGNES_API_KEY = os.getenv("AGNES_API_KEY")
        AGNES_API_URL = os.getenv("AGNES_API_URL", "https://apihub.agnes-ai.com/v1")
        AGNES_MODEL = os.getenv("AGNES_MODEL", "agnes-2.0-flash")

        if not AGNES_API_KEY:
            raise ValueError("AGNES_API_KEY 未配置，请在 .env 文件中设置")

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {AGNES_API_KEY}",
        }

        payload = {
            "model": AGNES_MODEL,
            "messages": messages,
            "temperature": 0.1,
            "max_tokens": 2000,
        }

        async with httpx.AsyncClient(timeout=120) as client:
            response = await client.post(
                AGNES_API_URL,
                headers=headers,
                json=payload,
            )
            response.raise_for_status()
            result = response.json()

        # 提取 AI 回复
        if "choices" in result and len(result["choices"]) > 0:
            tongue_analysis = result["choices"][0]["message"]["content"]
            return {"tongue_analysis": tongue_analysis}
        else:
            raise ValueError(f"AI API 返回格式异常: {result}")

    except Exception as e:
        print(f"舌象分析失败: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise



@app.post("/api/consultations")
async def create_consultation(
    consultation: ConsultationCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # Gather knowledge content if knowledge base specified
    knowledge_content = []
    custom_prompt = ""
    if consultation.knowledge_base_id:
        # Verify user has access to this knowledge base
        binding = db.query(KnowledgeBaseBinding).filter(
            KnowledgeBaseBinding.knowledge_base_id == consultation.knowledge_base_id,
            KnowledgeBaseBinding.user_id == current_user.id,
        ).first()
        if not binding and current_user.role != "super_admin":
            raise HTTPException(status_code=403, detail="您无权使用此知识库")

        docs = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.knowledge_base_id == consultation.knowledge_base_id
        ).all()
        knowledge_content = [d.parsed_content for d in docs if d.parsed_content]
        
        # Get custom prompt config from knowledge base
        if consultation.knowledge_base_id:
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == consultation.knowledge_base_id).first()
            if kb and kb.prompt_config:
                from datetime import datetime
                now = datetime.now().strftime("%Y年%m月%d日 %H:%M")
                custom_prompt = kb.prompt_config.replace("{current_time}", now)
            if kb and kb.negative_prompt:
                negative_prompt = kb.negative_prompt

    # Build messages
    messages = []
    has_tongue = bool(consultation.tongue_image_base64)

    if has_tongue:
        # Analyze tongue image first
        try:
            tongue_result = await analyze_tongue_image(consultation.tongue_image_base64, knowledge_content)
            tongue_analysis = tongue_result["tongue_analysis"]
        except Exception as e:
            tongue_analysis = f"舌象分析暂时不可用: {str(e)}"
    else:
        tongue_analysis = ""

    # Build combined question with tongue analysis + symptoms
    full_question = ""
    if tongue_analysis:
        full_question += f"【AI舌象辨析结果】\n{tongue_analysis}\n\n"
    if consultation.question.strip():
        full_question += f"【我的症状/问题】\n{consultation.question}"
    elif tongue_analysis:
        full_question += "请根据以上舌象辨析结果，给出综合辨证和养生方案。"
    else:
        full_question = "请描述您的症状或上传舌象图片。"

    messages.append({"role": "user", "content": full_question})

    try:
        ai_result = await call_agnes_ai(
            messages=messages,
            knowledge_content=knowledge_content,
            has_tongue_image=has_tongue,
            custom_prompt=custom_prompt,
            negative_prompt=negative_prompt,
        )
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="AI服务连接失败，请检查API地址配置")
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="AI服务响应超时，请稍后重试")
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"AI服务返回错误({e.response.status_code}): {e.response.text[:200]}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI服务暂时不可用: {str(e)}")

    # Extract symptoms from answer
    symptoms = []
    common_symptoms = ["头痛", "失眠", "乏力", "胸闷", "气短", "食欲不振", "腹胀", "便秘", "腹泻",
                       "眩晕", "耳鸣", "腰痛", "腿软", "怕冷", "怕热", "出汗", "口干", "口苦",
                       "心悸", "水肿"]
    for s in common_symptoms:
        if s in ai_result["answer"]:
            symptoms.append(s)

    # Save consultation
    record = Consultation(
        user_id=current_user.id,
        knowledge_base_id=consultation.knowledge_base_id,
        question=consultation.question,
        answer=ai_result["answer"],
        tongue_analysis=tongue_analysis,
        syndrome_analysis=ai_result["syndrome_analysis"],
        symptoms=json.dumps(symptoms, ensure_ascii=False),
    )
    db.add(record)

    # Save tongue image if provided
    if has_tongue and consultation.tongue_image_base64:
        img_data = base64.b64decode(consultation.tongue_image_base64)
        img_path = os.path.join(UPLOAD_DIR, f"tongue_{record.id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.jpg")
        with open(img_path, "wb") as f:
            f.write(img_data)
        tongue_img = TongueImage(
            consultation_id=record.id,
            user_id=current_user.id,
            image_path=img_path,
            analysis_result=tongue_analysis,
        )
        db.add(tongue_img)

    db.commit()
    db.refresh(record)

    return {
        "id": record.id,
        "reply": ai_result["answer"],
        "tongue_analysis": tongue_analysis,
        "syndrome_analysis": ai_result["syndrome_analysis"],
        "symptoms": symptoms,
    }


@app.post("/api/consultations/stream")
async def stream_consultation(
    consultation: ConsultationCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Stream AI consultation response using SSE"""
    # Gather knowledge content
    knowledge_content = []
    if consultation.knowledge_base_id:
        binding = db.query(KnowledgeBaseBinding).filter(
            KnowledgeBaseBinding.knowledge_base_id == consultation.knowledge_base_id,
            KnowledgeBaseBinding.user_id == current_user.id,
        ).first()
        if not binding and current_user.role != "super_admin":
            raise HTTPException(status_code=403, detail="您无权使用此知识库")
        docs = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.knowledge_base_id == consultation.knowledge_base_id
        ).all()
        knowledge_content = [d.parsed_content for d in docs if d.parsed_content]
        
        # Get custom prompt config from knowledge base
        custom_prompt = ""
        negative_prompt = ""
        if consultation.knowledge_base_id:
            kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == consultation.knowledge_base_id).first()
            if kb and kb.prompt_config:
                from datetime import datetime
                now = datetime.now().strftime("%Y年%m月%d日 %H:%M")
                custom_prompt = kb.prompt_config.replace("{current_time}", now)
            if kb and kb.negative_prompt:
                negative_prompt = kb.negative_prompt

    # Analyze tongue image first (not streamed, must complete before main answer)
    has_tongue = bool(consultation.tongue_image_base64)
    tongue_analysis = ""
    if has_tongue:
        try:
            tongue_result = await analyze_tongue_image(consultation.tongue_image_base64, knowledge_content)
            tongue_analysis = tongue_result["tongue_analysis"]
        except Exception as e:
            tongue_analysis = f"舌象分析暂时不可用: {str(e)}"

    # Build question
    full_question = ""
    if tongue_analysis:
        full_question += f"【AI舌象辨析结果】\n{tongue_analysis}\n\n"
    if consultation.question.strip():
        full_question += f"【我的症状/问题】\n{consultation.question}"
    elif tongue_analysis:
        full_question += "请根据以上舌象辨析结果，给出综合辨证和养生方案。"
    else:
        full_question = "请描述您的症状或上传舌象图片。"

    messages = [{"role": "user", "content": full_question}]

    async def event_generator():
        full_answer = ""
        try:
            # Send tongue analysis first if available
            if tongue_analysis:
                yield f"data: {json.dumps({'type': 'tongue', 'content': tongue_analysis}, ensure_ascii=False)}\n\n"

            # Stream AI response
            async for chunk in stream_agnes_ai(
                messages=messages,
                knowledge_content=knowledge_content,
                has_tongue_image=has_tongue,
                custom_prompt=custom_prompt,
                negative_prompt=negative_prompt,
            ):
                full_answer += chunk
                yield f"data: {json.dumps({'type': 'text', 'content': chunk}, ensure_ascii=False)}\n\n"

            # Extract symptoms
            symptoms = []
            common_symptoms = ["头痛", "失眠", "乏力", "胸闷", "气短", "食欲不振", "腹胀", "便秘", "腹泻",
                               "眩晕", "耳鸣", "腰痛", "腿软", "怕冷", "怕热", "出汗", "口干", "口苦",
                               "心悸", "水肿"]
            for s in common_symptoms:
                if s in full_answer:
                    symptoms.append(s)

            # Save to DB
            record = Consultation(
                user_id=current_user.id,
                knowledge_base_id=consultation.knowledge_base_id,
                question=consultation.question,
                answer=full_answer,
                tongue_analysis=tongue_analysis,
                syndrome_analysis=full_answer if "辨证" in full_answer else "",
                symptoms=json.dumps(symptoms, ensure_ascii=False),
            )
            db.add(record)

            # Save tongue image
            if has_tongue and consultation.tongue_image_base64:
                img_data = base64.b64decode(consultation.tongue_image_base64)
                tongue_img = TongueImage(
                    consultation_id=record.id,
                    image_data=img_data,
                    image_size=len(img_data),
                )
                db.add(tongue_img)

            db.commit()
            db.refresh(record)

            # Send done event with record id
            yield f"data: {json.dumps({'type': 'done', 'id': record.id, 'symptoms': symptoms}, ensure_ascii=False)}\n\n"

        except httpx.ConnectError:
            yield f"data: {json.dumps({'type': 'error', 'content': 'AI服务连接失败，请检查API地址配置'}, ensure_ascii=False)}\n\n"
        except httpx.TimeoutException:
            yield f"data: {json.dumps({'type': 'error', 'content': 'AI服务响应超时，请稍后重试'}, ensure_ascii=False)}\n\n"
        except httpx.HTTPStatusError as e:
            yield f"data: {json.dumps({'type': 'error', 'content': f'AI服务返回错误({e.response.status_code})'}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'content': f'AI服务暂时不可用: {str(e)}'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.get("/api/consultations")
async def list_consultations(
    knowledge_base_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(Consultation)

    # Role-based data access
    if current_user.role == "super_admin":
        pass  # Can see all consultations
    elif current_user.role == "admin":
        # 管理员：只看自己绑定知识库的问诊记录（不限区域）
        bindings = db.query(KnowledgeBaseBinding).filter(
            KnowledgeBaseBinding.user_id == current_user.id
        ).all()
        kb_ids = [b.knowledge_base_id for b in bindings]
        if kb_ids:
            query = query.filter(Consultation.knowledge_base_id.in_(kb_ids))
        else:
            query = query.filter(Consultation.user_id == current_user.id)
    elif current_user.role == "province_manager":
        # 省区总经理：只看该省区 + 绑定知识库的问诊记录
        bindings = db.query(KnowledgeBaseBinding).filter(
            KnowledgeBaseBinding.user_id == current_user.id
        ).all()
        kb_ids = [b.knowledge_base_id for b in bindings]
        # 关联User获取创建者的省份
        query = query.join(User, Consultation.user_id == User.id)
        if kb_ids:
            query = query.filter(
                Consultation.knowledge_base_id.in_(kb_ids),
                User.province == current_user.province
            )
        else:
            query = query.filter(
                User.province == current_user.province,
                Consultation.user_id == current_user.id
            )
    elif current_user.role == "region_manager":
        # 地区经理：只看该地区 + 绑定知识库的问诊记录
        bindings = db.query(KnowledgeBaseBinding).filter(
            KnowledgeBaseBinding.user_id == current_user.id
        ).all()
        kb_ids = [b.knowledge_base_id for b in bindings]
        query = query.join(User, Consultation.user_id == User.id)
        if kb_ids:
            query = query.filter(
                Consultation.knowledge_base_id.in_(kb_ids),
                User.province == current_user.province,
                User.region == current_user.region
            )
        else:
            query = query.filter(
                User.province == current_user.province,
                User.region == current_user.region,
                Consultation.user_id == current_user.id
            )
    elif current_user.role == "business_manager":
        # 业务经理：只看该地区 + 绑定知识库的问诊记录
        bindings = db.query(KnowledgeBaseBinding).filter(
            KnowledgeBaseBinding.user_id == current_user.id
        ).all()
        kb_ids = [b.knowledge_base_id for b in bindings]
        query = query.join(User, Consultation.user_id == User.id)
        if kb_ids:
            query = query.filter(
                Consultation.knowledge_base_id.in_(kb_ids),
                User.province == current_user.province,
                User.region == current_user.region
            )
        else:
            query = query.filter(Consultation.user_id == current_user.id)
    else:
        # 普通用户只能看到自己的问诊记录
        query = query.filter(Consultation.user_id == current_user.id)

    if knowledge_base_id:
        query = query.filter(Consultation.knowledge_base_id == knowledge_base_id)

    query = query.order_by(Consultation.created_at.desc())
    records = query.limit(200).all()
    return [_format_consultation(r, db) for r in records]


@app.get("/api/consultations/{consultation_id}")
async def get_consultation(
    consultation_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    record = db.query(Consultation).filter(Consultation.id == consultation_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="问诊记录不存在")

    if current_user.role not in ("super_admin", "admin"):
        subordinate_ids = _get_subordinate_ids(current_user, db)
        if record.user_id not in subordinate_ids + [current_user.id]:
            raise HTTPException(status_code=403, detail="权限不足")

    return _format_consultation(record, db)


@app.post("/api/local-cases")
async def create_local_case(
    data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    consultation_id = data.get("consultation_id")
    # Get consultation data if provided
    q, ans, tongue, syndrome, symptoms_data = "", "", "", "", "[]"
    if consultation_id:
        record = db.query(Consultation).filter(Consultation.id == consultation_id).first()
        if record:
            q = record.question
            ans = record.answer
            tongue = record.tongue_analysis or ""
            syndrome = record.syndrome_analysis or ""
            symptoms_data = record.symptoms or "[]"
            record.saved_locally = True
    # Use provided data as override
    case = LocalCase(
        user_id=current_user.id,
        consultation_id=consultation_id,
        patient_name=data.get("patient_name", ""),
        patient_gender=data.get("patient_gender", ""),
        patient_age=data.get("patient_age", ""),
        patient_phone=data.get("patient_phone", ""),
        patient_address=data.get("patient_address", ""),
        question=data.get("question", q),
        answer=data.get("answer", ans),
        tongue_analysis=data.get("tongue_analysis", tongue),
        syndrome_analysis=data.get("syndrome_analysis", syndrome),
        symptoms=symptoms_data,
    )
    db.add(case)
    db.commit()
    db.refresh(case)
    return {"id": case.id, "message": "案例已保存"}


@app.get("/api/local-cases")
async def list_local_cases(
    search: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(LocalCase).filter(LocalCase.user_id == current_user.id)
    if search:
        like = f"%{search}%"
        query = query.filter(
            (LocalCase.patient_name.like(like)) |
            (LocalCase.patient_phone.like(like))
        )
    cases = query.order_by(LocalCase.created_at.desc()).all()
    return [{
        "id": c.id,
        "patient_name": c.patient_name,
        "patient_gender": c.patient_gender,
        "patient_age": c.patient_age,
        "patient_phone": c.patient_phone,
        "patient_address": c.patient_address,
        "question": c.question,
        "reply": c.answer,
        "tongue_analysis": c.tongue_analysis,
        "syndrome_analysis": c.syndrome_analysis,
        "symptoms": json.loads(c.symptoms) if c.symptoms else [],
        "created_at": c.created_at.isoformat() if c.created_at else None,
    } for c in cases]


@app.get("/api/local-cases/{case_id}")
async def get_local_case(
    case_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    case = db.query(LocalCase).filter(LocalCase.id == case_id, LocalCase.user_id == current_user.id).first()
    if not case:
        raise HTTPException(status_code=404, detail="案例不存在")
    return {
        "id": case.id,
        "patient_name": case.patient_name,
        "patient_gender": case.patient_gender,
        "patient_age": case.patient_age,
        "patient_phone": case.patient_phone,
        "patient_address": case.patient_address,
        "question": case.question,
        "reply": case.answer,
        "tongue_analysis": case.tongue_analysis,
        "syndrome_analysis": case.syndrome_analysis,
        "symptoms": json.loads(case.symptoms) if case.symptoms else [],
        "created_at": case.created_at.isoformat() + "+08:00" if case.created_at else None,
    }


@app.delete("/api/local-cases/{case_id}")
async def delete_local_case(
    case_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    case = db.query(LocalCase).filter(LocalCase.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="案例不存在")
    # 超级管理员可删除所有，其他用户只能删除自己的
    if current_user.role != "超级管理员" and case.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="无权限删除此案例")
    db.delete(case)
    db.commit()
    return {"message": "案例已删除"}


# ==================== Analytics APIs ====================

@app.get("/api/analytics/overview")
async def analytics_overview(
    current_user: User = Depends(require_role("super_admin", "admin")),
    db: Session = Depends(get_db),
):
    total_users = db.query(User).count()
    total_consultations = db.query(Consultation).count()
    total_kbs = db.query(KnowledgeBase).count()
    total_docs = db.query(KnowledgeDocument).count()

    return {
        "total_users": total_users,
        "total_consultations": total_consultations,
        "total_knowledge_bases": total_kbs,
        "total_documents": total_docs,
    }


@app.get("/api/analytics/hot-symptoms")
async def hot_symptoms(
    limit: int = 20,
    current_user: User = Depends(require_role("super_admin", "admin", "province_manager", "region_manager", "business_manager")),
    db: Session = Depends(get_db),
):
    """Get top hot symptoms from consultation records"""
    query = db.query(Consultation.symptoms)

    if current_user.role not in ("super_admin", "admin"):
        subordinate_ids = _get_subordinate_ids(current_user, db)
        query = query.filter(Consultation.user_id.in_(subordinate_ids + [current_user.id]))

    records = query.all()

    symptom_count = {}
    for r in records:
        if r.symptoms:
            try:
                symptoms = json.loads(r.symptoms)
                for s in symptoms:
                    symptom_count[s] = symptom_count.get(s, 0) + 1
            except:
                pass

    sorted_symptoms = sorted(symptom_count.items(), key=lambda x: x[1], reverse=True)[:limit]
    return [{"symptom": s, "count": c} for s, c in sorted_symptoms]


@app.get("/api/analytics/consultation-stats")
async def consultation_stats(
    days: int = 30,
    current_user: User = Depends(require_role("super_admin", "admin", "province_manager", "region_manager", "business_manager")),
    db: Session = Depends(get_db),
):
    """Get consultation statistics by day"""
    from datetime import timedelta
    since = datetime.utcnow() - timedelta(days=days)

    query = db.query(
        func.date(Consultation.created_at).label("date"),
        func.count(Consultation.id).label("count"),
    ).filter(Consultation.created_at >= since)

    if current_user.role not in ("super_admin", "admin"):
        subordinate_ids = _get_subordinate_ids(current_user, db)
        query = query.filter(Consultation.user_id.in_(subordinate_ids + [current_user.id]))

    stats = query.group_by(func.date(Consultation.created_at)).order_by(func.date(Consultation.created_at)).all()

    return [{"date": str(s.date), "count": s.count} for s in stats]


# ==================== Helper Functions ====================

def _get_subordinate_ids(user: User, db: Session) -> List[int]:
    """Get all subordinate user IDs for a given user"""
    allowed_roles = SUBORDINATE_ROLES.get(user.role, [])
    query = db.query(User.id).filter(User.role.in_(allowed_roles), User.is_active == True)

    if user.province:
        query = query.filter(User.province == user.province)
    if user.region:
        query = query.filter(User.region == user.region)

    return [r[0] for r in query.all()]


def _format_consultation(record: Consultation, db: Session) -> dict:
    user = db.query(User).filter(User.id == record.user_id).first()
    kb = db.query(KnowledgeBase).filter(KnowledgeBase.id == record.knowledge_base_id).first() if record.knowledge_base_id else None
    symptoms = []
    if record.symptoms:
        try:
            symptoms = json.loads(record.symptoms)
        except:
            pass

    return {
        "id": record.id,
        "user_id": record.user_id,
        "username": user.username if user else "",
        "user_real_name": user.real_name if user else "",
        "knowledge_base_id": record.knowledge_base_id,
        "knowledge_base_name": kb.name if kb else "",
        "question": record.question,
        "reply": record.answer,
        "tongue_analysis": record.tongue_analysis,
        "syndrome_analysis": record.syndrome_analysis,
        "symptoms": symptoms,
        "saved_locally": record.saved_locally,
        "created_at": record.created_at.isoformat() if record.created_at else None,
    }


# ==================== Startup ====================

@app.on_event("startup")
async def startup():
    init_super_admin()
    # Auto-disable expired accounts
    from database import SessionLocal
    db = SessionLocal()
    try:
        now = datetime.utcnow()
        expired = db.query(User).filter(User.is_active == True, User.valid_until != None, User.valid_until < now).all()
        for u in expired:
            u.is_active = False
        if expired:
            db.commit()
            print(f"[INFO] Auto-disabled {len(expired)} expired account(s)")
    finally:
        db.close()


# ==================== Frontend Mount (must be last) ====================
# Mount frontend static files at root - this must be AFTER all API routes
from fastapi.responses import FileResponse

@app.get("/", include_in_schema=False)
async def serve_frontend_root():
    """Serve frontend index.html for root path"""
    index_file = FRONTEND_DIR / "index.html"
    if index_file.exists():
        return FileResponse(index_file, media_type="text/html")
    return {"error": "Frontend not found"}

# Mount static files (css, js)
static_dir = FRONTEND_DIR / "css"
if static_dir.exists():
    app.mount("/css", StaticFiles(directory=str(static_dir)), name="css")

static_dir = FRONTEND_DIR / "js"
if static_dir.exists():
    app.mount("/js", StaticFiles(directory=str(static_dir)), name="js")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

